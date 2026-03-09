"""
Embedding pipeline: file → chunks (unstructured) → embeddings (LiteLLM jsl-embed) → Qdrant + BM25.
"""

import re
import uuid
import httpx
from pathlib import Path
from typing import List, Dict, Any

from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

from ingestion.bm25_store import BM25Store

EMBED_DIM = 1024
COLLECTION = "jsl_docs"


def _infer_grade(filename: str) -> str:
    for grade in ["316L", "304", "430"]:
        if grade in filename:
            return grade
    return "unknown"


def _infer_doc_type(filename: str) -> str:
    fname = filename.lower()
    if "sop" in fname:
        return "sop"
    if "spec" in fname or "specification" in fname:
        return "spec"
    if "mainten" in fname or "manual" in fname:
        return "maintenance"
    if "checklist" in fname or "qc" in fname or "inspection" in fname:
        return "qc"
    if "sap" in fname:
        return "sap_process"
    if "safety" in fname or "environment" in fname or "compliance" in fname:
        return "policy"
    return "general"


def _extract_text_plain(file_path: str) -> List[Dict]:
    text = Path(file_path).read_text(errors="replace")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    doc_name = Path(file_path).name
    return [
        {
            "text": p,
            "metadata": {
                "doc_name": doc_name,
                "grade": _infer_grade(doc_name),
                "doc_type": _infer_doc_type(doc_name),
                "section": f"paragraph-{i+1}",
                "page": 1,
            },
        }
        for i, p in enumerate(paragraphs)
    ]


def _extract_pdf(file_path: str) -> List[Dict]:
    """Extract text from PDF using pdfminer (no OpenCV/libGL dependency)."""
    from pdfminer.high_level import extract_text_to_fp, extract_pages
    from pdfminer.layout import LTTextContainer
    import io

    doc_name = Path(file_path).name
    chunks = []

    # pdfminer layout gives one LTTextContainer per text box — treat each as a chunk,
    # grouping short consecutive boxes (headings) with the following body text.
    for page_num, page_layout in enumerate(extract_pages(file_path), start=1):
        current_section = f"page-{page_num}"
        pending = []  # accumulate short heading-like blocks

        for element in page_layout:
            if not isinstance(element, LTTextContainer):
                continue
            text = element.get_text().strip()
            if not text:
                continue

            # Short block (<60 chars, no punctuation end) → treat as section heading
            if len(text) < 60 and not text[-1] in ".,:;":
                if pending:
                    # flush previous accumulated text as a chunk
                    body = " ".join(pending)
                    if len(body) >= 20:
                        chunks.append({
                            "text": body,
                            "metadata": {
                                "doc_name": doc_name,
                                "grade": _infer_grade(doc_name),
                                "doc_type": _infer_doc_type(doc_name),
                                "section": current_section,
                                "page": page_num,
                            },
                        })
                    pending = []
                current_section = text[:80]
            else:
                pending.append(text)

        # flush remaining
        if pending:
            body = " ".join(pending)
            if len(body) >= 20:
                chunks.append({
                    "text": body,
                    "metadata": {
                        "doc_name": doc_name,
                        "grade": _infer_grade(doc_name),
                        "doc_type": _infer_doc_type(doc_name),
                        "section": current_section,
                        "page": page_num,
                    },
                })

    return chunks if chunks else _extract_text_plain(file_path)


def _extract_docx(file_path: str) -> List[Dict]:
    """Extract text from DOCX using python-docx (no OpenCV/libGL dependency)."""
    from docx import Document

    doc_name = Path(file_path).name
    doc = Document(file_path)
    chunks = []
    current_section = "Introduction"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        # Headings become section labels
        if para.style.name.startswith("Heading"):
            current_section = text[:80]
            continue
        chunks.append({
            "text": text,
            "metadata": {
                "doc_name": doc_name,
                "grade": _infer_grade(doc_name),
                "doc_type": _infer_doc_type(doc_name),
                "section": current_section,
                "page": 1,
            },
        })

    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if len(row_text) >= 20:
                chunks.append({
                    "text": row_text,
                    "metadata": {
                        "doc_name": doc_name,
                        "grade": _infer_grade(doc_name),
                        "doc_type": _infer_doc_type(doc_name),
                        "section": current_section,
                        "page": 1,
                    },
                })

    return chunks if chunks else _extract_text_plain(file_path)


def extract_chunks(file_path: str) -> List[Dict]:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        try:
            return _extract_pdf(file_path)
        except Exception:
            pass
    elif suffix in (".docx", ".doc"):
        try:
            return _extract_docx(file_path)
        except Exception:
            pass
    return _extract_text_plain(file_path)


def _ensure_collection(client: QdrantClient, collection: str) -> None:
    if not client.collection_exists(collection):
        client.create_collection(
            collection_name=collection,
            vectors_config=VectorParams(size=EMBED_DIM, distance=Distance.COSINE),
        )


def upsert_to_qdrant(
    client: QdrantClient,
    collection: str,
    chunks: List[Dict],
) -> None:
    _ensure_collection(client, collection)
    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=c["embedding"],
            payload={
                "text": c["text"],
                **c["metadata"],
            },
        )
        for c in chunks
    ]
    client.upsert(collection_name=collection, points=points)


class EmbeddingPipeline:
    def __init__(
        self,
        qdrant_client: QdrantClient,
        bm25_store: BM25Store,
        litellm_base_url: str,
        litellm_api_key: str,
        embed_model: str,
        collection: str = COLLECTION,
    ):
        self._qdrant = qdrant_client
        self._bm25 = bm25_store
        self._litellm_base_url = litellm_base_url.rstrip("/")
        self._litellm_api_key = litellm_api_key
        self._embed_model = embed_model
        self._collection = collection

    def _embed(self, texts: List[str]) -> List[List[float]]:
        all_embeddings = []
        batch_size = 32
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            resp = httpx.post(
                f"{self._litellm_base_url}/v1/embeddings",
                json={"model": self._embed_model, "input": batch},
                headers={"Authorization": f"Bearer {self._litellm_api_key}"},
                timeout=60.0,
            )
            resp.raise_for_status()
            data = resp.json()["data"]
            all_embeddings.extend([d["embedding"] for d in data])
        return all_embeddings

    def process(self, file_path: str) -> Dict[str, Any]:
        doc_id = str(uuid.uuid4())
        raw_chunks = extract_chunks(file_path)
        if not raw_chunks:
            return {"doc_id": doc_id, "chunks": 0, "status": "empty"}

        texts = [c["text"] for c in raw_chunks]
        embeddings = self._embed(texts)

        qdrant_chunks = [
            {**raw_chunks[i], "embedding": embeddings[i], "id": f"{doc_id}-{i}"}
            for i in range(len(raw_chunks))
        ]
        upsert_to_qdrant(self._qdrant, self._collection, qdrant_chunks)

        bm25_chunks = [
            {"id": f"{doc_id}-{i}", "text": raw_chunks[i]["text"]}
            for i in range(len(raw_chunks))
        ]
        self._bm25.add_chunks(bm25_chunks)
        self._bm25.save()

        return {"doc_id": doc_id, "chunks": len(raw_chunks), "status": "indexed"}
