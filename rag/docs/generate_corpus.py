"""
Generate 20 synthetic JSL stainless steel documents for the RAG demo corpus.
Produces PDF files (via reportlab) and Word files (via python-docx).

Usage:
    python rag/docs/generate_corpus.py
Output:
    rag/docs/corpus/ — 20 files ready for RAG indexing
"""

import os
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)
from reportlab.lib import colors
from docx import Document

OUTPUT_DIR = Path(__file__).parent / "corpus"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ── Chemical composition data (real-world values) ─────────────────────────
GRADE_SPECS = {
    "304": {
        "elements": [
            ("Carbon (C)", "max 0.08%"),
            ("Chromium (Cr)", "18.0–20.0%"),
            ("Nickel (Ni)", "8.0–10.5%"),
            ("Manganese (Mn)", "max 2.0%"),
            ("Silicon (Si)", "max 0.75%"),
            ("Phosphorus (P)", "max 0.045%"),
            ("Sulphur (S)", "max 0.030%"),
        ],
        "tensile_strength": "515 MPa min",
        "yield_strength": "205 MPa min",
        "hardness": "92 HRB max",
        "elongation": "40% min",
    },
    "316L": {
        "elements": [
            ("Carbon (C)", "max 0.03%"),
            ("Chromium (Cr)", "16.0–18.0%"),
            ("Nickel (Ni)", "10.0–14.0%"),
            ("Molybdenum (Mo)", "2.0–3.0%"),
            ("Manganese (Mn)", "max 2.0%"),
            ("Silicon (Si)", "max 0.75%"),
            ("Phosphorus (P)", "max 0.045%"),
            ("Sulphur (S)", "max 0.030%"),
        ],
        "tensile_strength": "485 MPa min",
        "yield_strength": "170 MPa min",
        "hardness": "95 HRB max",
        "elongation": "40% min",
    },
    "430": {
        "elements": [
            ("Carbon (C)", "max 0.12%"),
            ("Chromium (Cr)", "16.0–18.0%"),
            ("Manganese (Mn)", "max 1.0%"),
            ("Silicon (Si)", "max 1.0%"),
            ("Phosphorus (P)", "max 0.040%"),
            ("Sulphur (S)", "max 0.030%"),
        ],
        "tensile_strength": "450 MPa min",
        "yield_strength": "205 MPa min",
        "hardness": "88 HRB max",
        "elongation": "22% min",
    },
}

SURFACE_FINISH = {
    "304": "2B (cold rolled, annealed, pickled) — Ra ≤ 0.5 µm",
    "316L": "2B (cold rolled, annealed, pickled) — Ra ≤ 0.5 µm; BA option available",
    "430": "2B (cold rolled, annealed, pickled) — Ra ≤ 0.8 µm",
}


def _pdf_styles():
    styles = getSampleStyleSheet()
    heading1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=10,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
    )
    return heading1, body


def _make_table(data, col_widths=None):
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3a5c")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def generate_grade_spec(grade: str):
    spec = GRADE_SPECS[grade]
    path = OUTPUT_DIR / f"Grade_{grade}_Specification.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    h1, body = _pdf_styles()
    story = []

    story.append(Paragraph(f"Jindal Stainless Limited — Grade {grade} Specification", h1))
    story.append(Paragraph(f"Document No: JSL-SPEC-{grade}-2025 | Rev: 03 | Date: 2025-01-15", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Scope", h1))
    story.append(Paragraph(
        f"This specification defines the chemical composition, mechanical properties, "
        f"and surface finish requirements for JSL Grade {grade} stainless steel "
        f"cold-rolled coil and sheet.", body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. Applicable Standards", h1))
    story.append(Paragraph(
        "ASTM A240 / A240M | EN 10088-2 | IS 6911 | JIS G4304", body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("3.1 Product Forms", h1))
    story.append(Paragraph(
        "Cold-rolled coil, sheet, and strip. Thickness range: 0.3 mm – 6.0 mm. "
        "Width range: 600 mm – 1600 mm.", body))

    story.append(Paragraph("3.2 Chemical Composition", h1))
    table_data = [["Element", "Requirement"]] + list(spec["elements"])
    story.append(_make_table(table_data, col_widths=[8*cm, 8*cm]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("4. Mechanical Properties", h1))
    mech_data = [
        ["Property", "Requirement"],
        ["Tensile Strength (Rm)", spec["tensile_strength"]],
        ["0.2% Proof Strength (Rp0.2)", spec["yield_strength"]],
        ["Elongation (A50mm)", spec["elongation"]],
        ["Hardness", spec["hardness"]],
    ]
    story.append(_make_table(mech_data, col_widths=[8*cm, 8*cm]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("5. Surface Finish", h1))
    story.append(Paragraph(SURFACE_FINISH[grade], body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("6. Inspection & Testing", h1))
    story.append(Paragraph(
        "Each coil shall be subject to: (a) Chemical analysis per heat, "
        "(b) Mechanical testing per lot, (c) Visual surface inspection 100%, "
        "(d) Dimensional verification per coil.", body))

    doc.build(story)
    print(f"  ✓ {path.name}")


SOP_CONTENT = {
    "Pickling": {
        "purpose": "Remove oxide scale and restore corrosion resistance after annealing.",
        "steps": [
            ("1. Pre-rinse", "Rinse strip surface with demineralised water at 40–50°C for 30 seconds."),
            ("2. Acid bath entry", "Enter pickling section at line speed 15–25 m/min. Bath composition: HNO3 10–15%, HF 1–3%, temp 50–60°C."),
            ("3. Dwell time", "Minimum dwell: 45 seconds. Increase to 90s for heavy scale or Grade 430."),
            ("4. Rinse cascade", "3-stage counter-current rinse with pH ≤ 7.0 at final stage."),
            ("5. Passivation check", "Verify Ferroxyl test negative on 3 sample points per coil."),
            ("6. Dry-off", "Hot air knife at 80°C, ±5°C. Strip must be moisture-free before coiling."),
        ],
        "safety": "PPE mandatory: acid-resistant gloves, face shield, apron. Emergency shower within 10m.",
    },
    "Cold_Rolling": {
        "purpose": "Reduce strip thickness to target gauge with controlled surface finish.",
        "steps": [
            ("1. Roll gap setup", "Set initial roll gap via AGC system. Reference: setup sheet JSL-CR-SETUP-001."),
            ("2. Tension control", "Entry tension: 30–50 N/mm². Exit tension: 50–80 N/mm²."),
            ("3. Reduction ratio", "Max single-pass reduction: 35% for 304/316L, 30% for 430."),
            ("4. Roll force monitoring", "Alert if roll force exceeds 2500 T. Stop mill if >2800 T."),
            ("5. Surface roughness", "Measure Ra every 500m. Target: 0.3–0.5 µm for 2B finish."),
            ("6. Flatness control", "I-unit target ≤ 5 for coil grades. Increase tension if >8 I-unit."),
        ],
        "safety": "Lockout-tagout procedure mandatory for all roll changes. Minimum 2 operators required.",
    },
}

def generate_sop(sop_type: str, grade: str):
    sop = SOP_CONTENT[sop_type]
    safe_type = sop_type.replace("_", " ")
    path = OUTPUT_DIR / f"SOP_{sop_type}_{grade}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    h1, body = _pdf_styles()
    story = []

    story.append(Paragraph(
        f"Standard Operating Procedure — {safe_type} Line (Grade {grade})", h1))
    story.append(Paragraph(
        f"Doc No: JSL-SOP-{sop_type[:3].upper()}-{grade}-2025 | Rev: 02", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Purpose", h1))
    story.append(Paragraph(sop["purpose"], body))

    story.append(Paragraph("2. Scope", h1))
    story.append(Paragraph(
        f"Applies to all operators on the {safe_type} line processing Grade {grade} "
        f"stainless steel coil at JSL Hisar Works (Plant: JSL1, WERKS: JSL1).", body))

    story.append(Paragraph("3. Procedure Steps", h1))
    for step_name, step_desc in sop["steps"]:
        story.append(Paragraph(f"<b>{step_name}:</b> {step_desc}", body))

    story.append(Paragraph("4. Safety Requirements", h1))
    story.append(Paragraph(sop["safety"], body))

    story.append(Paragraph("5. Quality Records", h1))
    story.append(Paragraph(
        "Record all process parameters in SAP PM Work Order (AUFNR) linked to "
        "production order. Attach Ferroxyl / roughness test results as DMS attachments.", body))

    doc.build(story)
    print(f"  ✓ {path.name}")


MAINTENANCE_MANUALS = {
    "Annealing_Furnace": {
        "equipment_id": "EQ-ANN-001",
        "pm_schedule": [
            ("Daily", "Check burner flame pattern, thermocouple readings ±5°C accuracy"),
            ("Weekly", "Inspect refractory lining for cracks; clean recuperator fins"),
            ("Monthly", "Replace worn burner nozzles; calibrate O2 trim control"),
            ("Quarterly", "Full refractory inspection; replace damaged sections"),
            ("Annual", "Complete furnace reline; replace all thermocouples"),
        ],
        "bearing_replacement": "SKF 6316 bearings on hearth roll drives — replace every 8,000 hrs or if vibration > 4.5 mm/s RMS.",
        "lubrication": "Shell Omala S4 GX 220 for hearth roll gearboxes — change interval: 4,000 hrs.",
    },
    "Pickling_Line": {
        "equipment_id": "EQ-PCK-002",
        "pm_schedule": [
            ("Daily", "Check acid concentration: HNO3 10–15%, HF 1–3%. Titration every shift."),
            ("Weekly", "Inspect tank linings and joints for acid seepage"),
            ("Monthly", "Clean all rinse nozzles; check pH probe calibration"),
            ("Quarterly", "Replace all pump mechanical seals; inspect ventilation ducting"),
            ("Annual", "Full tank inspection and reline with acid-resistant coating"),
        ],
        "bearing_replacement": "FAG 22222-E1 spherical roller bearings on rinse mangle rolls — replace every 10,000 hrs.",
        "lubrication": "Castrol Tribol GR 100-2 PD for mangle roll bearings — relubricate every 500 hrs.",
    },
    "Rolling_Mill": {
        "equipment_id": "EQ-CRM-003",
        "pm_schedule": [
            ("Daily", "Check roll coolant flow rate ≥ 800 L/min; inspect strip tracking"),
            ("Weekly", "Measure work roll surface roughness; schedule regrind if Ra > 0.8 µm"),
            ("Monthly", "Inspect AGC hydraulics; check backup roll chock bearings"),
            ("Quarterly", "Full drive spindle inspection; check gear tooth wear"),
            ("Annual", "Mill stand alignment check; replace all chock seals"),
        ],
        "bearing_replacement": "NSK 239/800CAE4 for backup roll chocks — replace every 15,000 hrs or if clearance > 0.3 mm.",
        "lubrication": "Mobil Gear 600 XP 320 for main drive gearbox — change interval: 6,000 hrs.",
    },
}

def generate_maintenance_manual(equipment: str):
    manual = MAINTENANCE_MANUALS[equipment]
    safe_name = equipment.replace("_", " ")
    path = OUTPUT_DIR / f"Maintenance_Manual_{equipment}.docx"
    doc = Document()

    doc.add_heading(f"JSL Maintenance Manual — {safe_name}", 0)
    doc.add_paragraph(
        f"Equipment ID: {manual['equipment_id']} | Plant: JSL1 | "
        f"Doc: JSL-MM-{equipment[:3].upper()}-2025 | Rev: 01"
    )

    doc.add_heading("1. Preventive Maintenance Schedule", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Frequency"
    hdr[1].text = "Task"
    for freq, task in manual["pm_schedule"]:
        row = table.add_row().cells
        row[0].text = freq
        row[1].text = task

    doc.add_heading("2. Bearing Replacement", level=1)
    doc.add_paragraph(manual["bearing_replacement"])

    doc.add_heading("3. Lubrication", level=1)
    doc.add_paragraph(manual["lubrication"])

    doc.add_heading("4. SAP PM Integration", level=1)
    doc.add_paragraph(
        f"All maintenance activities must be logged against SAP PM Work Order "
        f"(AUFNR) for equipment {manual['equipment_id']}. Use transaction IW31 "
        f"to create corrective orders. Preventive orders auto-generated by "
        f"maintenance plan (WERKS: JSL1)."
    )

    doc.save(str(path))
    print(f"  ✓ {path.name}")


QC_CRITERIA = {
    "304": {
        "thickness_tol": "±0.05 mm for t ≤ 2mm; ±0.08 mm for t > 2mm",
        "width_tol": "+0 / -3 mm",
        "defect_limit": "No pits > 0.3 mm depth; no scratches > 0.2 mm depth; ≤ 2 edge marks per coil",
        "coil_weight_max": "20,000 kg",
    },
    "316L": {
        "thickness_tol": "±0.05 mm for t ≤ 2mm; ±0.08 mm for t > 2mm",
        "width_tol": "+0 / -3 mm",
        "defect_limit": "No pits > 0.2 mm depth; no scratches > 0.15 mm depth; zero edge marks for medical grade",
        "coil_weight_max": "18,000 kg",
    },
    "430": {
        "thickness_tol": "±0.06 mm for t ≤ 2mm; ±0.10 mm for t > 2mm",
        "width_tol": "+0 / -4 mm",
        "defect_limit": "No pits > 0.4 mm depth; no scratches > 0.25 mm depth; ≤ 3 edge marks per coil",
        "coil_weight_max": "22,000 kg",
    },
}

def generate_qc_checklist(grade: str):
    criteria = QC_CRITERIA[grade]
    path = OUTPUT_DIR / f"QC_Inspection_Checklist_{grade}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    h1, body = _pdf_styles()
    story = []

    story.append(Paragraph(f"QC Inspection Checklist — Grade {grade} Cold-Rolled Coil", h1))
    story.append(Paragraph(f"Form No: JSL-QC-{grade}-CHECK | Rev: 04 | Dept: Quality Assurance", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Dimensional Checks", h1))
    dim_data = [
        ["Parameter", "Acceptance Criteria"],
        ["Thickness tolerance", criteria["thickness_tol"]],
        ["Width tolerance", criteria["width_tol"]],
        ["Max coil weight", criteria["coil_weight_max"]],
    ]
    story.append(_make_table(dim_data, col_widths=[8*cm, 8*cm]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. Surface Defect Acceptance Criteria", h1))
    story.append(Paragraph(criteria["defect_limit"], body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. Visual Inspection Method", h1))
    story.append(Paragraph(
        "100% visual inspection under 1000 lux fluorescent lighting. "
        "Inspector traverses coil at 5 m/min. Mark defect locations with chalk. "
        "Photograph all reject defects. Log in SAP QM notification (QMEL).", body))

    story.append(Paragraph("4. Sampling for Mechanical Testing", h1))
    story.append(Paragraph(
        "One sample per lot (max 10 coils same heat). "
        "Tensile test per ASTM E8/E8M. Hardness per ASTM E18.", body))

    doc.build(story)
    print(f"  ✓ {path.name}")


def generate_sap_sop(process: str, content: dict):
    path = OUTPUT_DIR / f"SAP_{process}_SOP.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    h1, body = _pdf_styles()
    story = []
    story.append(Paragraph(content["title"], h1))
    story.append(Paragraph(content["docno"], body))
    story.append(Spacer(1, 12))
    for section_title, section_body in content["sections"]:
        story.append(Paragraph(section_title, h1))
        story.append(Paragraph(section_body, body))
        story.append(Spacer(1, 6))
    doc.build(story)
    print(f"  ✓ {path.name}")


SAP_SOPS = {
    "MM_Procurement": {
        "title": "SAP MM — Procurement SOP for Stainless Steel Raw Materials",
        "docno": "Doc: JSL-SAP-MM-001 | Module: MM | T-codes: ME21N, ME23N, MIGO",
        "sections": [
            ("1. Purchase Order Creation (ME21N)",
             "Raise PO against approved vendor. Key fields: EBELN (PO number), EBELP (line item), "
             "MATNR (material: e.g. STL-304-CR-2MM), MENGE (qty in MT), MEINS (UoM: MT), "
             "NETPR (net price INR/MT), WERKS (plant: JSL1)."),
            ("2. Goods Receipt (MIGO — 101)",
             "Post GR against PO. Movement type 101. "
             "System creates material document (MBLNR) and accounting document. "
             "Stock updates in unrestricted use (MARA-LABST)."),
            ("3. Invoice Verification (MIRO)",
             "Match vendor invoice to PO and GR. "
             "3-way match: PO price vs GR quantity vs invoice value. "
             "Post to AP if within tolerance ±2%."),
            ("4. Key SAP Tables",
             "MARA: Material master. EKKO: PO header. EKPO: PO item. "
             "MSEG: Material document segments. BSEG: Accounting document segments."),
        ],
    },
    "CO_Costing": {
        "title": "SAP CO — Production Order Costing Guide",
        "docno": "Doc: JSL-SAP-CO-001 | Module: CO-PC | T-codes: KKF6N, CO03, KKBC_ORD",
        "sections": [
            ("1. Production Order Cost Estimate",
             "Cost estimate calculated at order creation (AUFK). "
             "Planned cost = BOM components (RESB) + routing operations (AFVC). "
             "Key field: AUFNR (order number), MATNR (finished material), GAMNG (order qty)."),
            ("2. Actual Cost Posting",
             "Goods issue (MIGO-261) posts material cost to order. "
             "Confirmation (CO11N) posts activity costs (machine time, labour). "
             "Overhead applied by costing sheet."),
            ("3. Variance Analysis",
             "Run KKS2 to settle variances to cost centres. "
             "Price variance = actual vs standard price. "
             "Quantity variance = actual vs planned consumption."),
            ("4. Cost Variance Thresholds",
             "Alert threshold: variance > 5% of standard cost OR > INR 50,000 per order. "
             "Escalate to Plant Controller if variance > 10% for 3 consecutive months."),
        ],
    },
    "PM_Work_Order": {
        "title": "SAP PM — Maintenance Work Order SOP",
        "docno": "Doc: JSL-SAP-PM-001 | Module: PM | T-codes: IW31, IW32, IW41",
        "sections": [
            ("1. Create Corrective Work Order (IW31)",
             "Order type: PM01 (corrective). WERKS: JSL1. "
             "Enter equipment number (EQUNR) and functional location (TPLNR). "
             "Describe breakdown in short text. Assign to maintenance planner group."),
            ("2. Plan and Release",
             "Add operations (AFVC): work centre, duration, activity type. "
             "Assign materials from spare parts store (reservation). "
             "Release order (AUFNR status: REL) to enable time confirmation."),
            ("3. Time Confirmation (IW41)",
             "Actual hours entered per operation. "
             "Causes of failure: breakage (B1), wear (W2), corrosion (C3). "
             "System posts actual costs to order."),
            ("4. Technical Completion",
             "Set order to TECO when work complete. "
             "Settlement run (KO88) transfers costs to cost centre or asset. "
             "MTTR and MTBF updated automatically in equipment master."),
        ],
    },
}


def generate_policy_doc(filename: str, title: str, docno: str, sections: list):
    path = OUTPUT_DIR / filename
    doc_obj = SimpleDocTemplate(str(path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
    h1, body = _pdf_styles()
    story = []
    story.append(Paragraph(title, h1))
    story.append(Paragraph(docno, body))
    story.append(Spacer(1, 12))
    for sec_title, sec_body in sections:
        story.append(Paragraph(sec_title, h1))
        story.append(Paragraph(sec_body, body))
        story.append(Spacer(1, 6))
    doc_obj.build(story)
    print(f"  ✓ {path.name}")


POLICY_DOCS = [
    (
        "JSL_Safety_Standards_2025.pdf",
        "JSL Safety Standards — Industrial Operations 2025",
        "Doc: JSL-HSE-001 | Rev: 05 | Effective: 2025-01-01",
        [
            ("1. Chemical Handling Thresholds",
             "Hydrofluoric acid (HF): TLV-TWA 0.5 ppm. Nitric acid (HNO3): TLV-TWA 2 ppm. "
             "Mandatory continuous gas monitoring in pickling bay. Alarm at 50% TLV."),
            ("2. Machine Guarding",
             "All rotating equipment ≥ 50 rpm must have fixed guards. "
             "Pinch point distance: min 120 mm from nearest fixed structure. "
             "Guards removed only under LOTO procedure JSL-LOTO-STD-001."),
            ("3. Emergency Response",
             "Chemical spill: activate ERP, evacuate 20m radius, call ext. 999. "
             "Acid burn: flush with water 15 min minimum. Do not neutralise. "
             "AED locations: every production bay, max 3-minute walk."),
            ("4. Incident Reporting",
             "All incidents (near-miss and above) reported in SAP PM notification (QMEL) "
             "within 1 hour. LTI investigation completed within 24 hours. "
             "Monthly safety KPIs reported to plant head."),
        ],
    ),
    (
        "Environmental_Compliance_Report_2025.pdf",
        "JSL Environmental Compliance Report — FY 2025",
        "Doc: JSL-ENV-RPT-2025 | Prepared: Quality & Environment Dept",
        [
            ("1. Effluent Discharge Standards",
             "Pickling effluent: pH 6.0–9.0, Fluoride ≤ 10 mg/L, Total Chromium ≤ 2 mg/L. "
             "Monitored continuously by online analysers. Non-compliance triggers line stop."),
            ("2. Air Emission Limits",
             "Stack emissions from annealing furnace: NOx ≤ 400 mg/Nm³, SO2 ≤ 200 mg/Nm³. "
             "Continuous emission monitoring system (CEMS) installed per CPCB norms."),
            ("3. Waste Management",
             "Pickling sludge classified as hazardous waste (Schedule I). "
             "Disposal via CPCB-authorised vendor only. Manifest in Form 9 (HWM Rules 2016). "
             "Annual disposal: ~850 MT at JSL Hisar Works."),
            ("4. ISO 14001:2015 Compliance",
             "JSL Hisar Works certified ISO 14001:2015. "
             "Last surveillance audit: Oct 2025 — zero major non-conformances. "
             "Next recertification: Oct 2027."),
        ],
    ),
]


def main():
    print(f"Generating 20 JSL synthetic documents → {OUTPUT_DIR}/\n")

    print("Grade Specifications (PDF):")
    for grade in ["304", "316L", "430"]:
        generate_grade_spec(grade)

    print("\nPickling SOPs (PDF):")
    for grade in ["304", "316L", "430"]:
        generate_sop("Pickling", grade)

    print("\nCold Rolling SOPs (PDF):")
    for grade in ["304", "316L", "430"]:
        generate_sop("Cold_Rolling", grade)

    print("\nMaintenance Manuals (Word):")
    for equipment in ["Annealing_Furnace", "Pickling_Line", "Rolling_Mill"]:
        generate_maintenance_manual(equipment)

    print("\nQC Inspection Checklists (PDF):")
    for grade in ["304", "316L", "430"]:
        generate_qc_checklist(grade)

    print("\nSAP Process SOPs (PDF):")
    for process, content in SAP_SOPS.items():
        generate_sap_sop(process, content)

    print("\nPolicy Documents (PDF):")
    for filename, title, docno, sections in POLICY_DOCS:
        generate_policy_doc(filename, title, docno, sections)

    files = list(OUTPUT_DIR.iterdir())
    print(f"\n✅ Generated {len(files)} documents in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
